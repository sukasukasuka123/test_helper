"""
Agent 工具集（LangChain 版本，使用 @tool 装饰器，函数接收普通类型参数）
所有工具均用 @tool 注解定义，并保留 Pydantic 输入模型用于描述和验证
支持：姓名模糊匹配、批量处理、发送邮件、邮件附件下载、报名表索引与读取
"""
import json
import smtplib
import os
import imaplib
import email
import re
import sqlite3
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.header import decode_header
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()  # 加载项目根目录的 .env 文件

from langchain_core.tools import tool
from pydantic import BaseModel, Field, ValidationError

# ─────────────────────────────────────────────
# 163邮箱 IMAP 兼容性补丁
# 必须在模块加载时注册，否则 _simple_command('ID', ...) 会抛出 KeyError
# ─────────────────────────────────────────────
imaplib.Commands['ID'] = ('AUTH',)


# ─────────────────────────────────────────────
# Pydantic 参数 Schema（用于描述和验证）
# ─────────────────────────────────────────────

class LookupNameInput(BaseModel):
    """姓名查找工具输入参数"""
    name: str = Field(default="", description="面试者姓名，支持模糊匹配；传空字符串可列出所有人")


class AnalyzeInput(BaseModel):
    """面试者分析工具输入参数"""
    interviewee_ids: List[int] = Field(..., description="面试者 ID 列表，可传多个")


class ReportInput(BaseModel):
    """报告生成工具输入参数"""
    interviewee_ids: List[int] = Field(..., description="面试者 ID 列表")


class RecommendInput(BaseModel):
    """题目推荐工具输入参数"""
    interviewee_ids: List[int] = Field(..., description="面试者 ID 列表")
    num_questions: int = Field(default=3, description="每人推荐题目数量", ge=1, le=20)


class EmailRecipient(BaseModel):
    """邮件收件人信息"""
    interviewee_id: int = Field(..., description="面试者 ID（用于查询邮箱）")
    report_content: str = Field(..., description="邮件正文（通常是报告文本）")
    subject: Optional[str] = Field(default="您的面试报告", description="邮件主题")


class SendEmailInput(BaseModel):
    """发送邮件工具输入参数"""
    recipients: List[EmailRecipient] = Field(..., description="收件人列表，支持批量发送")


class GetDocInput(BaseModel):
    """获取邮件附件工具输入参数"""
    save_dir: str = Field(default="./attachments", description="附件保存目录路径")
    subject_filter: Optional[str] = Field(default=None, description="邮件主题关键词过滤，为空则获取所有含附件邮件")
    sender_filter: Optional[str] = Field(default=None, description="发件人邮箱过滤，为空则不限制")
    max_emails: int = Field(default=50, description="最多扫描邮件数量", ge=1, le=500)
    file_extensions: List[str] = Field(
        default=[".pdf", ".docx", ".doc", ".xlsx", ".xls", ".png", ".jpg", ".jpeg"],
        description="允许下载的附件扩展名列表"
    )


class WriteKeyInput(BaseModel):
    """建立报名表索引工具输入参数"""
    attachments_dir: str = Field(default="./attachments", description="附件所在目录路径")
    name_pattern: Optional[str] = Field(
        default=None,
        description="从文件名中提取姓名的正则表达式（含捕获组），如 r'报名表_(.+?)_\\d+'"
    )
    id_pattern: Optional[str] = Field(
        default=None,
        description="从文件名中提取学号的正则表达式（含捕获组），如 r'_(\\d{8,12})'"
    )


class ReadDocInput(BaseModel):
    """读取报名表内容工具输入参数"""
    file_path: str = Field(..., description="报名表文件的完整路径")
    extract_fields: List[str] = Field(
        default=["姓名", "学号", "专业", "年级", "联系方式", "邮箱", "意向部门"],
        description="需要从报名表中提取的字段名列表"
    )


class ReadKeyInput(BaseModel):
    """通过数据库查询索引内容工具输入参数"""
    name: Optional[str] = Field(default=None, description="按姓名模糊查询，为空则列出所有")
    student_id: Optional[str] = Field(default=None, description="按学号精确查询")
    limit: int = Field(default=50, description="最多返回条数", ge=1, le=500)


# ─────────────────────────────────────────────
# 数据库辅助：确保报名表索引表存在
# ─────────────────────────────────────────────

def _ensure_registration_table(db):
    """
    检查并创建 registration_index 表（若不存在）。
    表结构：id, name, student_id, file_path, file_name, created_at
    """
    db.execute("""
        CREATE TABLE IF NOT EXISTS registration_index (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            name        TEXT,
            student_id  TEXT,
            file_path   TEXT NOT NULL UNIQUE,
            file_name   TEXT,
            created_at  DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    """)


# ─────────────────────────────────────────────
# 1. 姓名查找工具
# ─────────────────────────────────────────────

def _create_lookup_tool(db):
    """工厂函数：创建姓名查找工具"""

    @tool(args_schema=LookupNameInput)
    def lookup_interviewees_by_name(name: str) -> str:
        """按姓名（支持模糊匹配）查找面试者，返回匹配的 ID 列表及基本信息。当用户提到人名时，必须先调用此工具获取 interviewee_id。"""
        name_val = name.strip() if name else ""
        if name_val:
            rows = db.fetchall(
                "SELECT id, name, email, phone FROM interviewee WHERE name LIKE ?",
                (f"%{name_val}%",)
            )
        else:
            rows = db.fetchall("SELECT id, name, email, phone FROM interviewee")

        if not rows:
            return f"未找到姓名包含「{name_val}」的面试者" if name_val else "暂无面试者记录"

        result = f"查找结果（共 {len(rows)} 人）:\n"
        for iid, iname, email, phone in rows:
            result += f"  - ID:{iid}  姓名:{iname}  邮箱:{email or '未填写'}  电话:{phone or '未填写'}\n"
        return result

    return lookup_interviewees_by_name


# ─────────────────────────────────────────────
# 2. 题库统计工具（无参数）
# ─────────────────────────────────────────────

def _create_question_stats_tool(db):
    """工厂函数：创建题库统计工具"""

    @tool
    def get_question_statistics() -> str:
        """获取题库统计信息，包括各类型、各难度的题目数量分布"""
        total = db.fetchall("SELECT COUNT(*) FROM question_bank")[0][0]

        type_stats = db.fetchall("""
            SELECT q_type, COUNT(*) as count
            FROM question_bank
            GROUP BY q_type
            ORDER BY count DESC
        """)
        diff_stats = db.fetchall("""
            SELECT difficulty, COUNT(*) as count
            FROM question_bank
            GROUP BY difficulty
            ORDER BY count DESC
        """)

        result = f"题库统计\n总题数: {total} 道\n\n类型分布:\n"
        for q_type, count in type_stats:
            result += f"  {q_type}: {count} 道\n"
        result += "\n难度分布:\n"
        for difficulty, count in diff_stats:
            result += f"  {difficulty}: {count} 道\n"
        return result

    return get_question_statistics


# ─────────────────────────────────────────────
# 3. 面试者分析工具（支持批量）
# ─────────────────────────────────────────────

def _create_analysis_tool(db):
    """工厂函数：创建面试者分析工具"""

    def _analyze_one(interviewee_id: int) -> str:
        info = db.fetchall(
            "SELECT name, email, created_at FROM interviewee WHERE id=?",
            (interviewee_id,)
        )
        if not info:
            return f"未找到面试者 ID={interviewee_id}"

        name, email, created_at = info[0]
        records = db.fetchall(
            "SELECT score, answer_snapshot FROM interview_record WHERE interviewee_id=?",
            (interviewee_id,)
        )

        if not records:
            return f"[{name}] 尚无答题记录"

        scores = [r[0] for r in records]
        avg_score = round(sum(scores) / len(scores), 2)

        type_scores: Dict[str, List] = {}
        for score, snap_json in records:
            snap = json.loads(snap_json)
            q_type = snap.get("type", "未知")
            type_scores.setdefault(q_type, []).append(score)

        rating = (
            "优秀" if avg_score >= 8 else
            "良好" if avg_score >= 6 else
            "及格" if avg_score >= 4 else "待提高"
        )

        result = (
            f"【{name}】(ID:{interviewee_id})\n"
            f"  邮箱: {email or '未填写'}  注册: {created_at}\n"
            f"  题数: {len(scores)}  总分: {sum(scores)}  均分: {avg_score}  "
            f"最高: {max(scores)}  最低: {min(scores)}\n"
            f"  各类型均分:\n"
        )
        for q_type, sc_list in type_scores.items():
            result += f"    {q_type}: {round(sum(sc_list) / len(sc_list), 2)} 分 ({len(sc_list)} 题)\n"
        result += f"  综合评级: {rating}\n"
        return result

    @tool(args_schema=AnalyzeInput)
    def analyze_interviewees(interviewee_ids: List[int]) -> str:
        """分析一个或多个面试者的答题表现（总分、均分、各类型得分、综合评级）。interviewee_ids 传入 ID 数组，支持批量分析。"""
        results = [_analyze_one(iid) for iid in interviewee_ids]
        return "\n\n" + ("=" * 60 + "\n").join(results)

    return analyze_interviewees


# ─────────────────────────────────────────────
# 4. 报告生成工具（支持批量）
# ─────────────────────────────────────────────

def _create_report_tool(db):
    """工厂函数：创建报告生成工具"""

    def _generate_one(interviewee_id: int) -> str:
        info = db.fetchall(
            "SELECT name, email, phone FROM interviewee WHERE id=?",
            (interviewee_id,)
        )
        if not info:
            return f"未找到面试者 ID={interviewee_id}"

        name, email, phone = info[0]
        records = db.fetchall("""
            SELECT question_id, score, answer_snapshot, created_at
            FROM interview_record
            WHERE interviewee_id = ?
            ORDER BY created_at
        """, (interviewee_id,))

        if not records:
            return f"[{name}] 无答题记录，无法生成报告"

        sep = "=" * 60
        report = f"{sep}\n{'面试报告':^56}\n{sep}\n"
        report += f"姓名: {name}  邮箱: {email or '未填写'}  电话: {phone or '未填写'}\n\n"
        report += "答题明细\n" + "-" * 60 + "\n"

        for idx, (q_id, score, snap_json, ans_time) in enumerate(records, 1):
            snap = json.loads(snap_json)
            report += (
                f"\n题目 {idx}  类型:{snap.get('type', '未知')}  "
                f"难度:{snap.get('difficulty', '未知')}  得分:{score}\n"
                f"  内容: {snap.get('content', '')[:60]}...\n"
                f"  时间: {ans_time}\n"
            )
            if snap.get("remark"):
                report += f"  备注: {snap['remark']}\n"

        scores = [r[1] for r in records]
        report += (
            f"\n{sep}\n统计分析\n"
            f"  题数:{len(scores)}  总分:{sum(scores)}  "
            f"均分:{round(sum(scores) / len(scores), 2)}  "
            f"最高:{max(scores)}  最低:{min(scores)}\n{sep}\n"
        )
        return report

    @tool(args_schema=ReportInput)
    def generate_reports(interviewee_ids: List[int]) -> str:
        """为一个或多个面试者生成详细面试报告（答题明细 + 统计分析）。返回报告文本，可配合 send_report_email 工具发送给面试者。"""
        reports = [_generate_one(iid) for iid in interviewee_ids]
        return "\n\n".join(reports)

    return generate_reports


# ─────────────────────────────────────────────
# 5. 题目推荐工具（支持批量）
# ─────────────────────────────────────────────

def _create_recommend_tool(db):
    """工厂函数：创建题目推荐工具"""

    def _recommend_one(interviewee_id: int, num_questions: int) -> str:
        info = db.fetchall(
            "SELECT name FROM interviewee WHERE id=?", (interviewee_id,)
        )
        if not info:
            return f"未找到面试者 ID={interviewee_id}"

        name = info[0][0]
        records = db.fetchall(
            "SELECT score, answer_snapshot FROM interview_record WHERE interviewee_id=?",
            (interviewee_id,)
        )

        if records:
            type_scores: Dict[str, List] = {}
            for score, snap_json in records:
                snap = json.loads(snap_json)
                q_type = snap.get("type", "未知")
                type_scores.setdefault(q_type, []).append(score)

            type_avg = {t: sum(sc) / len(sc) for t, sc in type_scores.items()}
            weak_type = min(type_avg, key=type_avg.get)
            weak_avg = type_avg[weak_type]

            recs = db.fetchall(
                "SELECT id, q_type, difficulty, content FROM question_bank WHERE q_type=? LIMIT ?",
                (weak_type, num_questions)
            )
            header = f"[{name}] 薄弱项「{weak_type}」(均分 {weak_avg:.2f})，推荐练习:\n"
        else:
            recs = db.fetchall(
                "SELECT id, q_type, difficulty, content FROM question_bank ORDER BY RANDOM() LIMIT ?",
                (num_questions,)
            )
            header = f"[{name}] 首次面试，随机推荐 {num_questions} 题:\n"

        if not recs:
            return f"[{name}] 题库暂无可推荐题目"

        result = header + "-" * 40 + "\n"
        for idx, (q_id, q_type, diff, content) in enumerate(recs, 1):
            result += f"  {idx}. [ID:{q_id}] {q_type} / {diff}\n     {content[:80]}...\n"
        return result

    @tool(args_schema=RecommendInput)
    def recommend_questions(interviewee_ids: List[int], num_questions: int = 3) -> str:
        """根据面试者历史表现，推荐合适题目（针对薄弱类型）。支持批量推荐。"""
        results = [_recommend_one(iid, num_questions) for iid in interviewee_ids]
        return "\n\n".join(results)

    return recommend_questions


# ─────────────────────────────────────────────
# 6. 发送邮件工具
# ─────────────────────────────────────────────

def _create_email_tool(db):
    """工厂函数：创建邮件发送工具"""

    smtp_config = {
        "host": os.getenv("SMTP_HOST", "smtp.163.com"),
        "port": int(os.getenv("SMTP_PORT", "465")),
        "user": os.getenv("SMTP_USER", ""),
        "pass": os.getenv("SMTP_AUID", ""),
        "from": os.getenv("SMTP_FROM", os.getenv("SMTP_USER", ""))
    }

    def _send_one(iid: int, subject: str, content: str) -> str:
        info = db.fetchall(
            "SELECT name, email FROM interviewee WHERE id=?", (iid,)
        )
        if not info:
            return f"❌ ID={iid} 未找到面试者"

        name, email = info[0]
        if not email:
            return f"❌ [{name}] 邮箱未填写，跳过发送"

        try:
            msg = MIMEMultipart()
            msg["From"] = smtp_config["from"]
            msg["To"] = email
            msg["Subject"] = subject
            msg.attach(MIMEText(content, "plain", "utf-8"))

            with smtplib.SMTP_SSL(smtp_config["host"], smtp_config["port"], timeout=15) as server:
                if smtp_config["user"]:
                    server.login(smtp_config["user"], smtp_config["pass"])
                server.sendmail(smtp_config["from"], email, msg.as_string())

            return f"✅ [{name}] 报告已发送至 {email}"
        except Exception as e:
            return f"❌ [{name}]({email}) 发送失败: {str(e)}"

    @tool(args_schema=SendEmailInput)
    def send_report_email(recipients: Union[List[Dict], List[EmailRecipient]]) -> str:
        """将面试报告通过邮件发送给指定面试者。recipients 为列表，每项包含 interviewee_id（用于获取邮箱）和 report_content（邮件正文）。支持批量发送。"""
        results = []
        for item in recipients:
            if isinstance(item, dict):
                try:
                    recipient = EmailRecipient.model_validate(item)
                except ValidationError as e:
                    results.append(f"❌ 收件人数据格式错误: {e}")
                    continue
            elif isinstance(item, EmailRecipient):
                recipient = item
            else:
                results.append(f"❌ 不支持的收件人类型: {type(item)}")
                continue

            results.append(_send_one(recipient.interviewee_id, recipient.subject, recipient.report_content))
        return "\n".join(results)

    return send_report_email


# ─────────────────────────────────────────────
# 7. 获取邮箱中邮件附件文件并下载到路径下的工具
# ─────────────────────────────────────────────

def _create_get_doc_tool():
    """工厂函数：创建邮件附件下载工具"""

    # IMAP 配置（从环境变量读取）
    imap_config = {
        "host": os.getenv("IMAP_HOST", "imap.163.com"),
        "port": int(os.getenv("IMAP_PORT", "993")),
        "user": os.getenv("IMAP_USER", os.getenv("SMTP_USER", "")),
        "pass": os.getenv("IMAP_PASS", os.getenv("SMTP_AUID", "")),
    }

    def _decode_str(s: str) -> str:
        """解码邮件头字段（处理 =?utf-8?...?= 格式）"""
        parts = decode_header(s)
        result = ""
        for part, charset in parts:
            if isinstance(part, bytes):
                result += part.decode(charset or "utf-8", errors="replace")
            else:
                result += part
        return result

    def _safe_filename(name: str) -> str:
        """将文件名中的非法字符替换为下划线"""
        return re.sub(r'[\\/:*?"<>|]', "_", name).strip()

    @tool(args_schema=GetDocInput)
    def get_email_attachments(
        save_dir: str = "./attachments",
        subject_filter: Optional[str] = None,
        sender_filter: Optional[str] = None,
        max_emails: int = 50,
        file_extensions: List[str] = None,
    ) -> str:
        """
        登录邮箱（IMAP），扫描收件箱中含附件的邮件，将符合条件的附件下载到指定目录。
        支持按主题关键词、发件人过滤，支持限定文件扩展名。
        下载完成后返回已保存的文件路径列表，可供后续建立索引使用。
        IMAP 账号信息从环境变量 IMAP_HOST / IMAP_PORT / IMAP_USER / IMAP_PASS 读取。
        """
        if file_extensions is None:
            file_extensions = [".pdf", ".docx", ".doc", ".xlsx", ".xls", ".png", ".jpg", ".jpeg"]

        # 规范化扩展名为小写
        allowed_ext = {ext.lower() for ext in file_extensions}

        # 创建保存目录
        save_path = Path(save_dir)
        save_path.mkdir(parents=True, exist_ok=True)

        if not imap_config["user"] or not imap_config["pass"]:
            return "❌ IMAP 账号未配置，请检查环境变量 IMAP_USER / IMAP_PASS"

        downloaded: List[str] = []
        skipped: int = 0
        errors: List[str] = []

        try:
            # ── 连接 & 认证 ───────────────────────────────────────
            mail = imaplib.IMAP4_SSL(imap_config["host"], imap_config["port"])
            mail.login(imap_config["user"], imap_config["pass"])

            # ── 163 安全策略：登录后必须先发 ID 命令声明客户端身份 ──
            # 不发此命令直接 SELECT 会触发「Unsafe Login」被拦截
            _user_prefix = imap_config["user"].split("@")[0]
            _id_args = (
                "name", "PythonIMAPClient",
                "version", "1.0.0",
                "vendor", "internal-tool",
                "contact", _user_prefix,
            )
            _id_str = '"' + '" "'.join(_id_args) + '"'
            mail._simple_command("ID", f"({_id_str})")

            # 等待风控策略生效（163要求）
            import time
            time.sleep(2)

            mail.select("INBOX")

            # 搜索邮件
            search_criteria = "ALL"
            if sender_filter:
                search_criteria = f'FROM "{sender_filter}"'

            _, msg_ids_raw = mail.search(None, search_criteria)
            msg_ids = msg_ids_raw[0].split()

            # 取最近 max_emails 封（倒序，优先最新）
            msg_ids = msg_ids[-max_emails:][::-1]

            for mid in msg_ids:
                try:
                    _, msg_data = mail.fetch(mid, "(RFC822)")
                    raw = msg_data[0][1]
                    msg = email.message_from_bytes(raw)

                    # 主题过滤
                    subject = _decode_str(msg.get("Subject", ""))
                    if subject_filter and subject_filter not in subject:
                        continue

                    # 遍历附件
                    for part in msg.walk():
                        content_disposition = part.get("Content-Disposition", "")
                        if "attachment" not in content_disposition:
                            continue

                        raw_filename = part.get_filename()
                        if not raw_filename:
                            continue

                        filename = _safe_filename(_decode_str(raw_filename))
                        ext = Path(filename).suffix.lower()
                        if ext not in allowed_ext:
                            skipped += 1
                            continue

                        # 避免重名：若已存在则跳过
                        target = save_path / filename
                        if target.exists():
                            skipped += 1
                            continue

                        # 写入文件
                        payload = part.get_payload(decode=True)
                        if payload:
                            target.write_bytes(payload)
                            downloaded.append(str(target.resolve()))

                except Exception as e:
                    errors.append(f"处理邮件 {mid} 出错: {e}")
                    continue

            mail.logout()

        except Exception as e:
            return f"❌ 连接邮箱失败: {e}"

        lines = [
            f"📥 附件下载完成，保存目录: {save_path.resolve()}",
            f"  成功下载: {len(downloaded)} 个文件",
            f"  跳过（重复或类型不符）: {skipped} 个",
        ]
        if errors:
            lines.append(f"  错误: {len(errors)} 条")
            lines.extend(f"    - {e}" for e in errors[:5])
        if downloaded:
            lines.append("  已下载文件列表:")
            for fp in downloaded:
                lines.append(f"    • {fp}")

        return "\n".join(lines)

    return get_email_attachments


# ─────────────────────────────────────────────
# 8. 给附件建立报名表索引到数据库的工具（姓名-学号-报名表路径）
# ─────────────────────────────────────────────

def _create_write_key_tool(db):
    """工厂函数：创建报名表索引写入工具"""

    @tool(args_schema=WriteKeyInput)
    def write_registration_index(
        attachments_dir: str = "./attachments",
        name_pattern: Optional[str] = None,
        id_pattern: Optional[str] = None,
    ) -> str:
        """
        扫描指定目录下的所有附件文件，尝试从文件名中提取姓名和学号，
        并将【姓名 - 学号 - 文件路径】写入数据库 registration_index 表。
        若数据库中该表不存在，则自动创建；若文件路径已存在则跳过（避免重复索引）。
        name_pattern / id_pattern 为可选的正则表达式，用于从文件名提取信息；
        若不提供则尝试通用规则（下划线分隔）。
        """
        # ── 确保数据库表存在 ──────────────────────────
        _ensure_registration_table(db)

        dir_path = Path(attachments_dir)
        if not dir_path.exists():
            return f"❌ 目录不存在: {attachments_dir}"

        # 编译正则（可选）
        name_re = re.compile(name_pattern) if name_pattern else None
        id_re = re.compile(id_pattern) if id_pattern else None

        def _extract_name_id(filename: str):
            """
            从文件名中提取姓名和学号。
            优先使用用户提供的正则；否则尝试以下通用规则：
              - 格式1: 姓名_学号_*.ext   →  首段为姓名，第二段为学号
              - 格式2: 学号_姓名_*.ext   →  首段全数字则视为学号
              - 格式3: 姓名（学号）*.ext
              - 兜底: 文件名作为姓名，学号置空
            """
            stem = Path(filename).stem
            extracted_name, extracted_id = None, None

            if name_re:
                m = name_re.search(stem)
                extracted_name = m.group(1) if m else None
            if id_re:
                m = id_re.search(stem)
                extracted_id = m.group(1) if m else None

            # 若均未提供正则，使用通用规则
            if not name_re and not id_re:
                # 规则：括号内学号
                bracket_m = re.search(r'[（(](\d{6,12})[）)]', stem)
                if bracket_m:
                    extracted_id = bracket_m.group(1)
                    extracted_name = stem[:bracket_m.start()].strip("_- ")
                else:
                    parts = re.split(r'[_\-\s]+', stem)
                    if len(parts) >= 2:
                        if re.fullmatch(r'\d{6,12}', parts[0]):
                            extracted_id, extracted_name = parts[0], parts[1]
                        elif re.fullmatch(r'\d{6,12}', parts[1]):
                            extracted_name, extracted_id = parts[0], parts[1]
                        else:
                            extracted_name = parts[0]
                    else:
                        extracted_name = stem

            return extracted_name, extracted_id

        inserted, skipped, errors = 0, 0, []

        # 递归扫描目录中所有文件
        all_files = [f for f in dir_path.rglob("*") if f.is_file()]

        for fpath in all_files:
            abs_path = str(fpath.resolve())
            try:
                name_val, id_val = _extract_name_id(fpath.name)

                # 检查是否已存在（file_path UNIQUE 约束）
                existing = db.fetchall(
                    "SELECT id FROM registration_index WHERE file_path=?", (abs_path,)
                )
                if existing:
                    skipped += 1
                    continue

                db.execute(
                    "INSERT INTO registration_index (name, student_id, file_path, file_name) VALUES (?, ?, ?, ?)",
                    (name_val, id_val, abs_path, fpath.name)
                )
                inserted += 1
            except Exception as e:
                errors.append(f"{fpath.name}: {e}")

        lines = [
            f"📋 报名表索引完成（目录: {dir_path.resolve()}）",
            f"  新增索引: {inserted} 条",
            f"  已跳过（重复）: {skipped} 条",
            f"  总文件数: {len(all_files)} 个",
        ]
        if errors:
            lines.append(f"  错误: {len(errors)} 条")
            lines.extend(f"    - {e}" for e in errors[:5])
        return "\n".join(lines)

    return write_registration_index


# ─────────────────────────────────────────────
# 9. 读取报名表内容并返回对应内容格式的工具（增强版）
# ─────────────────────────────────────────────

def _create_read_doc_tool():
    """工厂函数：创建报名表内容读取工具（含注入检测与字段对齐）"""

    # ── 安全检测配置 ──────────────────────────
    # 基于实际攻击样例（黄子然 - 面试报名表.xlsx）增强的特征库
    INJECTION_PATTERNS = [
        r"扮演.*?角色",  # 检测角色扮演请求
        r"从现在开始",  # 检测指令覆盖时间状语
        r"忽略.*?指令",  # 检测忽略指令
        r"系统指令",  # 检测系统指令关键词
        r"假设你是",  # 检测假设性身份
        r"you are.*?now",  # 英文注入常见模式
        r"ignore.*?previous",  # 英文忽略历史
        r"%\s*[你您]",  # 检测特殊符号开头的中文指令 (针对样例 %你)
        r"%.*?(?:扮演 | 假设 | 指令 | 忽略)",  # 增强：% 开头后跟指令动词
        r"###.*?指令",  # 检测 Markdown 分隔符注入
        r"猫娘|persona|system prompt",  # 针对样例的具体高风险关键词
        r"接下来.*?过程中",  # 检测持续性指令覆盖
    ]
    COMPILED_PATTERNS = [re.compile(p, re.IGNORECASE) for p in INJECTION_PATTERNS]

    def _detect_injection(text: str) -> List[Dict[str, Any]]:
        """检测文本中是否包含提示词注入特征，返回详细警告信息"""
        warnings = []
        for i, pattern in enumerate(COMPILED_PATTERNS):
            matches = pattern.findall(text)
            if matches:
                # 记录匹配到的具体片段，便于定位
                warnings.append({
                    "pattern": INJECTION_PATTERNS[i],
                    "matches": matches[:3],  # 只保留前 3 个匹配项避免过长
                    "severity": "HIGH" if "猫娘" in matches[0] or "%" in INJECTION_PATTERNS[i] else "MEDIUM"
                })
        return warnings

    def _read_pdf(file_path: str) -> str:
        """提取 PDF 文本"""
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        except ImportError:
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            except ImportError:
                return "[错误] 未安装 pdfplumber 或 pypdf，无法读取 PDF"

    def _read_docx(file_path: str) -> str:
        """提取 DOCX 文本"""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_texts.append(" | ".join(cells))
            return "\n".join(paragraphs + table_texts)
        except ImportError:
            return "[错误] 未安装 python-docx，无法读取 DOCX"

    def _read_xlsx(file_path: str) -> str:
        """提取 XLSX 文本"""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        # 使用 | 分隔单元格，便于后续解析 Key-Value 对
                        lines.append(" | ".join(cells))
            return "\n".join(lines)
        except ImportError:
            return "[错误] 未安装 openpyxl，无法读取 XLSX"

    def _extract_fields(text: str, fields: List[str]) -> Dict[str, str]:
        """从文本中按字段名提取值（优化了分隔符处理）"""
        result = {}
        for field in fields:
            # 优化正则：允许字段名后跟 ':' 或 '|' 作为分隔符，不再将 '|' 视为终止符
            # 适配 _read_xlsx 生成的 "字段名： | 值" 或 "字段名：值" 格式
            pattern = rf'{re.escape(field)}\s*[：:]\s*\|?\s*(.+?)(?=\n|$)'
            m = re.search(pattern, text)
            result[field] = m.group(1).strip() if m else ""
        return result

    @tool(args_schema=ReadDocInput)
    def read_registration_doc(file_path: str, extract_fields: List[str] = None) -> str:
        """
        读取指定路径下的报名表文件，提取结构化字段，并进行安全注入检测。
        针对实验室报名表模板优化，特别关注“申请理由”等高风险字段。
        返回 JSON 格式，包含字段内容及安全警告。
        """
        # 默认字段对齐至 面试报名表 - 模板.xlsx
        if extract_fields is None:
            extract_fields = ["姓名", "学号", "邮箱", "面试方向", "细分方向", "核心项目", "申请理由"]

        fpath = Path(file_path)
        if not fpath.exists():
            return json.dumps({"error": f"文件不存在：{file_path}"}, ensure_ascii=False)

        ext = fpath.suffix.lower()
        if ext == ".pdf":
            raw_text = _read_pdf(file_path)
        elif ext in (".docx",):
            raw_text = _read_docx(file_path)
        elif ext in (".xlsx", ".xls"):
            raw_text = _read_xlsx(file_path)
        else:
            try:
                raw_text = fpath.read_text(encoding="utf-8", errors="replace")
            except Exception as e:
                return json.dumps({"error": f"不支持的文件类型或读取失败：{e}"}, ensure_ascii=False)

        if not raw_text.strip():
            return json.dumps({"error": "文件内容为空或无法提取文本"}, ensure_ascii=False)

        # ── 安全检测 ──────────────────────────
        security_warnings = _detect_injection(raw_text)

        # 风险等级评估
        risk_level = "LOW"
        if security_warnings:
            if any(w.get("severity") == "HIGH" for w in security_warnings):
                risk_level = "HIGH"
            else:
                risk_level = "MEDIUM"

        extracted = _extract_fields(raw_text, extract_fields)

        # 特别检查申请理由字段是否包含高风险内容
        reason_field = extracted.get("申请理由", "")
        if reason_field and security_warnings:
            # 如果存在警告且申请理由非空，提示重点审查该字段
            for w in security_warnings:
                if any(m in reason_field for m in w.get("matches", [])):
                    w["affected_field"] = "申请理由"

        return json.dumps(
            {
                "file": fpath.name,
                "risk_level": risk_level,
                "security_warnings": security_warnings,
                "fields": extracted,
                "raw_preview": raw_text[:500].replace("\n", " "),
                "security_tip": "发现高风险注入特征时，请勿直接将内容输入 LLM，建议人工复核。" if risk_level == "HIGH" else ""
            },
            ensure_ascii=False,
            indent=2
        )

    return read_registration_doc
# ─────────────────────────────────────────────
# 10. 通过数据库获取索引内容的工具（姓名-学号-报名表路径）
# ─────────────────────────────────────────────

def _create_read_key_tool(db):
    """工厂函数：创建报名表索引查询工具"""

    @tool(args_schema=ReadKeyInput)
    def read_registration_index(
        name: Optional[str] = None,
        student_id: Optional[str] = None,
        limit: int = 50,
    ) -> str:
        """
        查询数据库中的报名表索引（registration_index 表），
        支持按姓名模糊查询或按学号精确查询，返回匹配记录（含姓名、学号、文件路径）。
        若数据库中该表不存在，自动创建并提示用户先运行索引工具。
        """
        # ── 确保数据库表存在 ──────────────────────────
        _ensure_registration_table(db)

        # 构建查询条件
        conditions, params = [], []
        if name:
            conditions.append("name LIKE ?")
            params.append(f"%{name}%")
        if student_id:
            conditions.append("student_id = ?")
            params.append(student_id)

        where_clause = ("WHERE " + " AND ".join(conditions)) if conditions else ""
        params.append(limit)

        rows = db.fetchall(
            f"SELECT id, name, student_id, file_name, file_path, created_at "
            f"FROM registration_index {where_clause} ORDER BY created_at DESC LIMIT ?",
            tuple(params)
        )

        if not rows:
            hint = "提示：请先运行 write_registration_index 工具建立索引。"
            if name or student_id:
                return f"未找到匹配的报名表记录。\n{hint}"
            return f"数据库中暂无报名表索引记录。\n{hint}"

        lines = [f"查询结果（共 {len(rows)} 条）："]
        for rid, rname, rid_num, fname, fpath, created in rows:
            lines.append(
                f"  [{rid}] 姓名:{rname or '未知'}  学号:{rid_num or '未知'}  "
                f"文件:{fname}  路径:{fpath}  录入时间:{created}"
            )
        return "\n".join(lines)

    return read_registration_index


# ─────────────────────────────────────────────
# 工具注册入口
# ─────────────────────────────────────────────

def get_default_tools(db) -> List:
    """
    获取所有默认工具的 LangChain Tool 列表（使用 @tool 装饰器创建）

    用法：
        tools = get_default_tools(db)
        agent.register_tools(tools)
    """
    return [
        _create_lookup_tool(db),
        _create_question_stats_tool(db),
        _create_analysis_tool(db),
        _create_report_tool(db),
        _create_recommend_tool(db),
        _create_email_tool(db),
        _create_get_doc_tool(),           # 工具7：邮件附件下载（无需 db）
        _create_write_key_tool(db),       # 工具8：建立报名表索引
        _create_read_doc_tool(),          # 工具9：读取报名表内容（无需 db）
        _create_read_key_tool(db),        # 工具10：查询报名表索引
    ]


def register_default_tools(agent, db):
    """
    向后兼容函数：直接注册到 Agent 实例
    （内部调用 get_default_tools + agent.register_tools）
    """
    tools = get_default_tools(db)
    agent.register_tools(tools)
    print(f"[AgentTools] 已注册 {len(tools)} 个 LangChain 工具")